"""Renders the agent's structured content into a polished .docx file."""
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ACCENT = RGBColor(0x1F, 0x4E, 0x79)


def build_docx(title: str, document_type: str, sections: list[dict], assumptions: list[str], output_path: str):
    doc = Document()

    # base style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title block
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.color.rgb = ACCENT

    meta = doc.add_paragraph()
    meta.add_run(f"Document type: {document_type.replace('_', ' ').title()}").italic = True
    meta.add_run(f"   |   Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}").italic = True

    doc.add_paragraph()  # spacer

    # Assumptions callout, if any (transparency about agent decisions)
    if assumptions:
        doc.add_heading("Agent Assumptions", level=2)
        for a in assumptions:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(a)
        doc.add_paragraph()

    # Sections
    for sec in sections:
        doc.add_heading(sec["heading"], level=1)
        if sec.get("format") == "bullets" and sec.get("bullets"):
            for b in sec["bullets"]:
                doc.add_paragraph(b, style="List Bullet")
        elif sec.get("table"):
            rows = sec["table"]
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Light Grid Accent 1"
            for r, row in enumerate(rows):
                for c, val in enumerate(row):
                    table.cell(r, c).text = str(val)
        else:
            doc.add_paragraph(sec.get("body", ""))

    # Footer
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.text = "Generated autonomously by AI Document Agent"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(output_path)
    return output_path
